import os
import re
import json
import tempfile
import subprocess
import sys
import unittest

sys.path.insert(0, os.path.abspath("src"))
from docx import Document

from heuristics import infer_q27_baseline, infer_q28_safety
from intake_parser import (
    Q13_HEADER,
    Q15_HEADER,
    Q16_HEADER,
    Q17_HEADER,
    Q18_HEADER,
    Q19_HEADER,
    Q20_HEADER,
    Q21_HEADER,
    Q22_HEADER,
    Q23_HEADER,
    Q24_HEADER,
    Q25_HEADER,
    Q27_HEADER,
    Q28_HEADER,
    build_engine_answers,
)
from keyword_extractor import extract_keywords, variability_bar
from report_payload_builder import build_executive_blocks, build_hybrid_plan, build_report_payload
from content_library.external_snapshot import build_external_snapshot
from content_library.executive_sentences import build_executive_sentences
from scoring.bars import build_bar, bar_5, render_bar
from scoring.entity import infer_entity_type
from scoring.exposure import score_round_display
from scoring.external_tags import compute_external_structural_tags
from scoring.safetynet import safety_net_internal
from scoring.urgency import urgency_multiplier
from scoring.external_snapshot import infer_snapshot_type


class PipelineTests(unittest.TestCase):
    Q6_DECISION_KEY = "Section 2. Main Career Decision\n\n6. What is the main career decision you are considering right now? \n(Paragraph)"
    ALLOWED_COMPLETENESS = {
        "High (qualitative signals available).",
        "Moderate (some qualitative inputs limited).",
        "Partial (key qualitative signals missing).",
    }

    def _base_row(self):
        return {
            Q13_HEADER: "mixed outlook",
            Q15_HEADER: "some differentiation",
            Q16_HEADER: "medium confidence",
            Q17_HEADER: "mixed signal with uncertainty",
            Q18_HEADER: "some changes",
            Q19_HEADER: "maybe support",
            Q20_HEADER: "somewhat clear",
            Q21_HEADER: "mixed pace",
            Q22_HEADER: "some exposure but manageable",
            Q23_HEADER: "balance",
            Q24_HEADER: "somewhat comfortable",
            Q25_HEADER: "gradually through testing",
            Q27_HEADER: "very stressful but going to go to work",
            Q28_HEADER: "I have almost no safety net if this fails.",
        }

    def test_keyword_extraction_en(self):
        text = "There is restructuring risk, but I have a sponsor and enough savings."
        extracted = extract_keywords(text, "en")
        self.assertIn("restructuring cycle", extracted["market"])
        self.assertIn("sponsor support", extracted["internal"])
        self.assertIn("savings buffer", extracted["personal"])

    def test_keyword_extraction_ko(self):
        text = "구조조정 가능성이 있고 스폰서가 있으며 비상금도 있습니다."
        extracted = extract_keywords(text, "ko")
        self.assertIn("구조조정 압력", extracted["market"])
        self.assertIn("스폰서 지원", extracted["internal"])
        self.assertIn("비상자금 버퍼", extracted["personal"])

    def test_q27_inference_downside_upside_neutral(self):
        downside = "layoff risk, burnout, and conflict are likely with uncertainty"
        upside = "promotion, growth, and opportunity with visibility and raise"
        neutral = "some changes are possible but mostly similar outcomes"
        self.assertEqual(infer_q27_baseline(downside, "en"), "downside_heavy")
        self.assertEqual(infer_q27_baseline(upside, "en"), "upside_room")
        self.assertEqual(infer_q27_baseline(neutral, "en"), "neutral")

    def test_q28_inference_none_strong_some(self):
        none_case = "no plan b and no savings"
        strong_case = "I have severance and spouse income"
        some_case = "I am not sure but may have a possible option"
        self.assertEqual(infer_q28_safety(none_case, "en"), "none")
        self.assertEqual(infer_q28_safety(strong_case, "en"), "strong")
        self.assertEqual(infer_q28_safety(some_case, "en"), "some")

    def test_variability_bar_mapping(self):
        self.assertEqual(variability_bar(0), "▯▯▯▯")
        self.assertEqual(variability_bar(1), "▮▯▯▯")
        self.assertEqual(variability_bar(2), "▮▮▯▯")
        self.assertEqual(variability_bar(3), "▮▮▮▯")

    def test_missing_required_column_raises(self):
        row = {
            Q15_HEADER: "some differentiation",
            Q16_HEADER: "medium",
            Q17_HEADER: "mixed",
            Q18_HEADER: "some",
            Q19_HEADER: "maybe",
            Q20_HEADER: "somewhat clear",
            Q21_HEADER: "mixed",
            Q22_HEADER: "some exposure",
            Q23_HEADER: "balance",
            Q24_HEADER: "somewhat",
            Q25_HEADER: "gradually",
            Q27_HEADER: "neutral statement",
            Q28_HEADER: "plan b exists",
        }
        with self.assertRaises(ValueError):
            build_engine_answers(row, "en")

    def test_exec_blocks_generated_non_empty(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in (
            "Exec_Block_1_Title",
            "Exec_Block_1_Line",
            "Exec_Block_2_Title",
            "Exec_Block_2_Line",
            "Exec_Block_3_Title",
            "Exec_Block_3_Line",
            "Exec_Block_4_Title",
            "Exec_Block_4_Line",
        ):
            self.assertIn(key, payload)
            self.assertTrue(str(payload[key]).strip())

    def test_data_completeness_note_exists_and_allowed(self):
        payload = build_report_payload(self._base_row(), "en")
        self.assertIn("Data_Completeness_Note", payload)
        self.assertTrue(str(payload["Data_Completeness_Note"]).strip())
        self.assertIn(payload["Data_Completeness_Note"], self.ALLOWED_COMPLETENESS)

    def test_entity_inference_program(self):
        result = infer_entity_type("Considering a PhD application this year", "Option includes doctoral program")
        self.assertEqual(result["entity_type"], "program")
        self.assertEqual(result["entity_label"], "program")

    def test_entity_inference_venture(self):
        result = infer_entity_type("Considering startup founding path", "Option includes building a business venture")
        self.assertEqual(result["entity_type"], "venture")
        self.assertEqual(result["entity_label"], "venture")

    def test_diagnosis_defaults_generated_non_empty(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in (
            "D1_Core_Read",
            "D1_Implication",
            "D2_Core_Read",
            "D2_Implication",
            "D3_Core_Read",
            "D3_Implication",
            "D4_Core_Read",
            "D4_Implication",
            "D5_Core_Read",
            "D5_Implication",
        ):
            self.assertIn(key, payload)
            self.assertTrue(str(payload[key]).strip())

    def test_fifwm_scores_and_notes_generated(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in (
            "F_Formal_Score",
            "F_Informal_Score",
            "F_Framework_Score",
            "F_Workflow_Score",
            "F_MarketPolicy_Score",
        ):
            self.assertIn(key, payload)
            self.assertEqual(int(payload[key]), 1)
        for key in (
            "F_Formal_Note",
            "F_Informal_Note",
            "F_Framework_Note",
            "F_Workflow_Note",
            "F_MarketPolicy_Note",
        ):
            self.assertIn(key, payload)
            self.assertTrue(str(payload[key]).strip())

    def test_score_round_display(self):
        self.assertEqual(score_round_display(0.49), 0)
        self.assertEqual(score_round_display(0.5), 1)
        self.assertEqual(score_round_display(1.49), 1)
        self.assertEqual(score_round_display(1.5), 2)

    def test_bar_mapping_internal_fraction(self):
        self.assertEqual(bar_5(0), "▯▯▯")
        self.assertEqual(bar_5(1), "▮▯▯")
        self.assertEqual(bar_5(2), "▮▮▮")

    def test_bar_zero(self):
        self.assertEqual(render_bar(0.0), "▯▯▯")

    def test_bar_full(self):
        self.assertEqual(render_bar(2.0), "▮▮▮")

    def test_bar_mid(self):
        self.assertEqual(render_bar(1.0), "▮▯▯")

    def test_build_bar_mapping(self):
        self.assertEqual(build_bar(0), "▯▯▯")
        self.assertEqual(build_bar(1), "▮▯▯")
        self.assertEqual(build_bar(2), "▮▮▮")

    def test_urgency_multiplier_values(self):
        self.assertEqual(urgency_multiplier("Within 3 months"), 1.15)
        self.assertEqual(urgency_multiplier("Within 6 months"), 1.05)
        self.assertEqual(urgency_multiplier("No immediate timeline"), 1.0)

    def test_safetynet_none_vs_lt6(self):
        self.assertEqual(safety_net_internal("", "no safety net"), 0.0)
        self.assertEqual(safety_net_internal("", "<6 months of savings"), 0.5)

    def test_executive_blocks_no_placeholder_strings(self):
        blocks = build_executive_blocks(
            {
                "verdict": "Conditional Go",
                "exposure_score": 1,
                "market_score": 1,
                "safety_net_level": "low",
                "motivation_type": "Growth",
            }
        )
        for key in ("verdict", "outlook", "risk", "personal"):
            self.assertTrue(blocks[key].strip())
            self.assertNotIn("{{", blocks[key])
            self.assertNotEqual(blocks[key].strip(), "One-line structural summary.")

    def test_executive_sentences_deterministic(self):
        context = {
            "verdict": "Conditional Go",
            "market_score": 1,
            "risk_score": 1,
            "fit_score": 1,
            "exposure_score": 1,
            "safety_net_level": "moderate",
            "case_id": "case-123",
            "Full_Name": "Case User",
        }
        a = build_executive_sentences(context)
        b = build_executive_sentences(context)
        self.assertEqual(a, b)
        self.assertTrue(a["Exec_Block_1_Line"].strip())
        self.assertTrue(a["Exec_Block_2_Line"].strip())
        self.assertTrue(a["Exec_Block_3_Line"].strip())
        self.assertTrue(a["Exec_Block_4_Line"].strip())
        for key in ("Exec_Block_1_Line", "Exec_Block_2_Line", "Exec_Block_3_Line", "Exec_Block_4_Line"):
            self.assertNotIn("{{", a[key])
            self.assertNotEqual(a[key].strip(), "One-line structural summary.")

    def test_hybrid_plan_always_three(self):
        plan = build_hybrid_plan({"verdict": "Conditional Go"})
        self.assertEqual(len(plan), 3)
        for item in plan:
            self.assertIn("title", item)
            self.assertIn("context", item)
            self.assertIn("objective", item)
            self.assertIn("timeline", item)
            self.assertIn("success_signal", item)
            self.assertIn("abort_trigger", item)

    def test_payload_has_external_inputs_stable_strings(self):
        payload = build_report_payload(self._base_row(), "en")
        self.assertIn("external_inputs", payload)
        ext = payload["external_inputs"]
        self.assertEqual(
            set(ext.keys()),
            {"decision_text", "options_text", "baseline_text", "safetynet_text", "location_text"},
        )
        for value in ext.values():
            self.assertIsInstance(value, str)

    def test_external_tags_all_medium_when_external_inputs_empty(self):
        tags = compute_external_structural_tags(
            {
                "external_inputs": {
                    "decision_text": "",
                    "options_text": "",
                    "baseline_text": "",
                    "safetynet_text": "",
                    "location_text": "",
                }
            }
        )
        self._assert_external_tag_shape(tags)
        self.assertEqual(
            tags,
            {
                "irreversibility": "medium",
                "income_compression": "medium",
                "credential_dependency": "medium",
                "competition_density": "medium",
                "institutional_volatility": "medium",
                "mobility_load": "medium",
            },
        )

    def test_external_tags_korean_trigger_high_irreversibility_and_income(self):
        tags = compute_external_structural_tags(
            {
                "external_inputs": {
                    "decision_text": "박사 유학을 위해 퇴사 고려",
                    "options_text": "",
                    "baseline_text": "예적금 1년 버틸 수 있으나 소득 단절 우려",
                    "safetynet_text": "",
                    "location_text": "",
                }
            }
        )
        self._assert_external_tag_shape(tags)
        self.assertEqual(tags["irreversibility"], "high")
        self.assertEqual(tags["income_compression"], "high")

    def test_external_snapshot_keys_exist_non_empty(self):
        payload = build_report_payload(self._base_row(), "en")
        keys = (
            "External_Block_1_Title",
            "External_Block_1_Line",
            "External_Block_2_Title",
            "External_Block_2_Line",
            "External_Block_3_Title",
            "External_Block_3_Line",
            "External_Block_4_Title",
            "External_Block_4_Line",
        )
        for key in keys:
            self.assertIn(key, payload)
            self.assertTrue(str(payload[key]).strip())

    def test_external_snapshot_line_length_caps(self):
        snap = build_external_snapshot(
            {
                "decision_text": "phd full-time relocate move abroad with visa",
                "options_text": "highly competitive top schools",
                "baseline_text": "restructuring and policy shock concerns",
                "safetynet_text": "no safety net under 6 months",
                "location_text": "seoul",
            },
            {
                "irreversibility": "high",
                "income_compression": "high",
                "credential_dependency": "high",
                "competition_density": "high",
                "institutional_volatility": "high",
                "mobility_load": "high",
            },
        )
        for i in range(1, 5):
            self.assertLessEqual(len(snap[f"External_Block_{i}_Title"]), 32)
            self.assertLessEqual(len(snap[f"External_Block_{i}_Line"]), 140)

    def test_external_snapshot_no_placeholders(self):
        snap = build_external_snapshot(
            {"decision_text": "", "options_text": "", "baseline_text": "", "safetynet_text": "", "location_text": ""},
            {"irreversibility": "medium", "income_compression": "medium", "credential_dependency": "medium", "competition_density": "medium", "institutional_volatility": "medium", "mobility_load": "medium"},
        )
        for i in range(1, 5):
            line = snap[f"External_Block_{i}_Line"]
            self.assertNotIn("TBD", line)
            self.assertNotIn("(1 line)", line)
            self.assertNotIn("Core read", line)

    def _assert_external_tag_shape(self, tags):
        self.assertEqual(
            set(tags.keys()),
            {
                "irreversibility",
                "income_compression",
                "credential_dependency",
                "competition_density",
                "institutional_volatility",
                "mobility_load",
            },
        )
        for value in tags.values():
            self.assertIn(value, {"low", "medium", "high"})

    def test_external_tags_phd_crossborder_high_irreversible_credential(self):
        tags = compute_external_structural_tags(
            {
                "external_inputs": {
                    "decision_text": "Mid-career full-time PhD transition with cross-border move and visa dependency.",
                    "options_text": "Target includes top schools and highly competitive academic track.",
                    "baseline_text": "Baseline appears uncertain with mixed directional evidence.",
                    "safetynet_text": "Almost no safety net, less than 6 months buffer.",
                    "location_text": "Korea, Seoul",
                }
            }
        )
        self._assert_external_tag_shape(tags)
        self.assertEqual(tags["irreversibility"], "high")
        self.assertEqual(tags["income_compression"], "high")
        self.assertEqual(tags["credential_dependency"], "high")
        self.assertEqual(tags["competition_density"], "high")
        self.assertEqual(tags["institutional_volatility"], "medium")
        self.assertEqual(tags["mobility_load"], "high")

    def test_external_tags_internal_move_low_irreversible_low_mobility(self):
        tags = compute_external_structural_tags(
            {
                "external_inputs": {
                    "decision_text": "Internal transfer in same company and same city using short pilot rotation.",
                    "options_text": "Niche internal promotion with unique profile fit and no new credential.",
                    "baseline_text": "Environment remains stable and predictable with no major changes.",
                    "safetynet_text": "12 months emergency fund and no relocation.",
                    "location_text": "USA, Seattle",
                }
            }
        )
        self._assert_external_tag_shape(tags)
        self.assertEqual(tags["irreversibility"], "low")
        self.assertEqual(tags["income_compression"], "low")
        self.assertEqual(tags["credential_dependency"], "low")
        self.assertEqual(tags["competition_density"], "low")
        self.assertEqual(tags["institutional_volatility"], "low")
        self.assertEqual(tags["mobility_load"], "low")

    def test_external_tags_income_compression_driven_by_low_buffer(self):
        tags = compute_external_structural_tags(
            {
                "external_inputs": {
                    "decision_text": "Industry move into adjacent function.",
                    "options_text": "External role in competitive but non-elite market.",
                    "baseline_text": "Recent layoffs and reorg indicate elevated institutional volatility.",
                    "safetynet_text": "Savings buffer around 8 months and limited fallback support.",
                    "location_text": "",
                }
            }
        )
        self._assert_external_tag_shape(tags)
        self.assertEqual(tags["institutional_volatility"], "high")
        self.assertEqual(tags["income_compression"], "medium")
        self.assertEqual(tags["irreversibility"], "medium")
        self.assertEqual(tags["mobility_load"], "medium")

    def test_external_snapshot_type_industry_default(self):
        t = infer_snapshot_type({"decision_text": "considering a promotion or role change"})
        self.assertEqual(t, "industry_transition")

    def test_external_snapshot_title_academic(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                self.Q6_DECISION_KEY: "Full-time PhD application decision",
            },
            "en",
        )
        self.assertEqual(payload["External_Snapshot_Title"], "External Snapshot — Academic Transition")

    def test_external_snapshot_title_industry(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                self.Q6_DECISION_KEY: "Considering a promotion or role change",
            },
            "en",
        )
        self.assertEqual(payload["External_Snapshot_Title"], "External Snapshot — Industry Transition")

    def test_mobility_block_category_switch_for_academic_transition(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                self.Q6_DECISION_KEY: "Move from corporate leadership to a full-time PhD in academia.",
            },
            "en",
        )
        self.assertEqual(payload["Mobility_Type"], "Category Switch")
        self.assertTrue(str(payload["Mobility_Reading"]).strip())
        self.assertTrue(str(payload["Mobility_Implication"]).strip())

    def test_mobility_block_internal_expansion_for_promotion(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                self.Q6_DECISION_KEY: "Internal promotion with bigger scope in the same company.",
            },
            "en",
        )
        self.assertEqual(payload["Mobility_Type"], "Internal Expansion")
        self.assertTrue(str(payload["Mobility_Reading"]).strip())
        self.assertTrue(str(payload["Mobility_Implication"]).strip())

    def test_external_snapshot_payload_and_merge_no_placeholders(self):
        payload = build_report_payload(self._base_row(), "en")
        self.assertIn("External_Snapshot_Type", payload)
        self.assertIn("External_Snapshot_Title", payload)
        self.assertTrue(str(payload["External_Snapshot_Title"]).strip())
        self.assertIn("External_Snapshot", payload)
        snapshot_lines = [ln.strip() for ln in str(payload["External_Snapshot"]).splitlines() if ln.strip()]
        self.assertEqual(len(snapshot_lines), 4)
        self.assertEqual(len(set(snapshot_lines)), 4)

        with tempfile.TemporaryDirectory() as td:
            template_path = os.path.join(td, "template.docx")
            payload_path = os.path.join(td, "payload.json")
            out_path = os.path.join(td, "merged.docx")

            doc = Document()
            doc.add_paragraph("External Snapshot Type: {{External_Snapshot_Type}}")
            doc.add_paragraph("{{External_Snapshot}}")
            doc.save(template_path)

            with open(payload_path, "w", encoding="utf-8") as fh:
                json.dump(payload, fh, ensure_ascii=False)

            subprocess.run(
                [
                    sys.executable,
                    "src/merge_docx.py",
                    "--template",
                    template_path,
                    "--payload",
                    payload_path,
                    "--out",
                    out_path,
                ],
                check=True,
            )

            merged = Document(out_path)
            full_text = "\n".join(p.text for p in merged.paragraphs)
            self.assertNotRegex(full_text, re.compile(r"\{\{External_Snapshot[^}]*\}\}"))

    def test_internal_structural_snapshot_exists(self):
        payload = build_report_payload(self._base_row(), "en")
        self.assertIn("Internal_Structural_Snapshot", payload)

    def test_internal_structural_snapshot_non_empty_and_not_placeholder(self):
        payload = build_report_payload(self._base_row(), "en")
        text = str(payload.get("Internal_Structural_Snapshot", "")).strip()
        self.assertTrue(text)
        self.assertNotIn("{{", text)
        self.assertNotIn("TBD", text.upper())

    def test_value_structure_fields_exist(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in (
            "Value_Current_Path",
            "Value_Transition_Path",
            "Value_Structure_Reading",
            "Value_Structure_Implication",
        ):
            self.assertIn(key, payload)

    def test_value_structure_fields_non_empty(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in (
            "Value_Current_Path",
            "Value_Transition_Path",
            "Value_Structure_Reading",
            "Value_Structure_Implication",
        ):
            self.assertIsInstance(payload[key], str)
            self.assertTrue(payload[key].strip())

    def test_value_structure_phd_case_mentions_key_tension_terms(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                self.Q6_DECISION_KEY: "Considering quitting current role for a full-time PhD transition.",
                "8. What career options are you currently considering?\nPlease briefly describe each option in 2–3 sentences.\nYou may list them as Option 1, Option 2, etc.": "PhD study with relocation versus staying in current corporate track.",
                Q28_HEADER: "I have almost no safety net and low savings buffer.",
            },
            "en",
        )
        combined = (
            f"{payload.get('Value_Structure_Reading', '')} {payload.get('Value_Structure_Implication', '')}"
        ).lower()
        self.assertTrue(
            any(term in combined for term in ("stability", "compression", "optionality", "transition cost"))
        )

    def test_temperament_fields_exist_and_non_empty(self):
        payload = build_report_payload(self._base_row(), "en")
        for key in ("Temperament_Profile", "Temperament_Reading", "Temperament_Implication"):
            self.assertIn(key, payload)
            self.assertIsInstance(payload[key], str)
            self.assertTrue(payload[key].strip())

    def test_temperament_enterprising_case(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                Q23_HEADER: "maximize upside with visible risk acceptance",
                Q24_HEADER: "comfortable with taking visible risks",
                Q25_HEADER: "all-in commitment once chosen",
            },
            "en",
        )
        self.assertEqual(payload["Temperament_Profile"], "Enterprising leaning")

    def test_temperament_defensive_case(self):
        payload = build_report_payload(
            {
                **self._base_row(),
                Q23_HEADER: "protecting the downside first",
                Q24_HEADER: "not comfortable with visible risks",
                Q25_HEADER: "small, reversible testing steps gradually",
            },
            "en",
        )
        self.assertEqual(payload["Temperament_Profile"], "Defensive leaning")


if __name__ == "__main__":
    unittest.main()
