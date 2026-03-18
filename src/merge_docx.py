import argparse
import json
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--payload", default="output/report_payload_latest.json")
    ap.add_argument("--out", default="output/AMC_Report_Latest.docx")
    ap.add_argument("--strict-undeclared", action="store_true")
    args = ap.parse_args()

    with open(args.payload, "r", encoding="utf-8") as f:
        ctx = json.load(f)
    if not isinstance(ctx, dict):
        raise SystemExit("Payload root must be a JSON object for docxtpl context.")

    tpl = DocxTemplate(args.template)
    undeclared = sorted(tpl.get_undeclared_template_variables(context=ctx))
    if undeclared:
        print("UNDECLARED_TEMPLATE_VARIABLES:", len(undeclared))
        for key in undeclared:
            print("  -", key)
        if args.strict_undeclared:
            raise SystemExit("Template has undeclared variables for provided context.")

    tpl.render(ctx)

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(out_path))

    # Post-merge leftover placeholder audit.
    doc = Document(str(out_path))
    leftover = set()
    for p in doc.paragraphs:
        text = p.text or ""
        if "{{" in text and "}}" in text:
            leftover.add(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text or ""
                    if "{{" in text and "}}" in text:
                        leftover.add(text)
    if leftover:
        print("LEFTOVER_PLACEHOLDER_LINES:", len(leftover))
        for line in sorted(leftover):
            print("  -", line)
    print("WROTE:", out_path)

if __name__ == "__main__":
    main()
